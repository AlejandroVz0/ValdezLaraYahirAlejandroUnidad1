from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
CAPTURES_DIR = DOCS_DIR / "capturas"
EVIDENCE_PATH = DOCS_DIR / "evidencias" / "backend-evidencia.json"
OUTPUT_PATH = DOCS_DIR / "U1-SH-DWP-ValdezLaraYahirAlejandro.docx"

REPO_URL = "https://github.com/AlejandroVz0/ValdezLaraYahirAlejandroUnidad1"


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def style_normal(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 12,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, *, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    document.add_paragraph("")


def build_cover(document: Document) -> None:
    add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        document,
        "Ingenieria en Desarrollo y Gestion de Software",
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        "Desarrollo Web Profesional",
        bold=True,
        size=15,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        document,
        "Unidad 1 - Proyecto Web",
        bold=True,
        size=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        "Ferreteria Valdez",
        bold=True,
        size=22,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        "Sitio web con frontend Stitch, catalogo, formularios y backend Node.js",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_table(
        document,
        ["Dato", "Informacion"],
        [
            ["Elaborado por", "Yahir Alejandro Valdez Lara"],
            ["Matricula", "23040027"],
            ["Grupo", "8 IDGS B"],
            ["Docente", "Gabriel Alejandro Reyes Morales"],
            ["Repositorio", REPO_URL],
        ],
    )

    add_paragraph(
        document,
        "Ramos Arizpe, Coah. | 08 jun 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    document.add_page_break()


def build_body(document: Document) -> None:
    evidence = json.loads(EVIDENCE_PATH.read_text(encoding="utf-8"))

    add_heading(document, "Introduccion")
    add_paragraph(
        document,
        "En este proyecto se desarrollo el sitio web de Ferreteria Valdez, que es una ferreteria ficticia inspirada en negocios reales de Ramos Arizpe, Coahuila. La idea principal fue crear una pagina que se viera moderna, ordenada y util para los clientes, pero que al mismo tiempo cumpliera con todos los requisitos que se pidieron en la materia de Desarrollo Web Profesional. Durante la realizacion del trabajo se fueron integrando varias secciones importantes como el inicio, el catalogo de productos, el area de ayuda, el contacto, el mapa del sitio y la pagina de error 404. Tambien se agregaron formularios de registro, inicio de sesion y recuperacion de contrasena para que el sitio se sintiera mas completo y funcional.\n\nAdemas de la parte visual, este proyecto tambien incluye un backend basico en Node.js para guardar informacion en archivos JSON, como usuarios, mensajes y productos. Eso hace que la pagina no sea solo de presentacion, sino que tambien tenga funciones reales para practicar mejor lo aprendido. En general, este trabajo me ayudo a entender mejor como se relacionan el frontend y el backend dentro de un mismo proyecto, y tambien me permitio reforzar temas como validaciones, rutas, estructura de carpetas, estilo responsivo y documentacion final para una entrega formal.",
    )

    add_heading(document, "Tecnologias utilizadas")
    add_table(
        document,
        ["Tecnologia", "Uso"],
        [
            ["React + Vite", "Frontend y empaquetado del proyecto"],
            ["Tailwind CSS", "Estilos y diseno responsivo"],
            ["Node.js", "Servidor y endpoints API"],
            ["JSON", "Almacenamiento de productos, usuarios y mensajes"],
            ["GitHub", "Entrega y publicacion del proyecto"],
        ],
    )

    add_heading(document, "Requisitos cumplidos")
    add_table(
        document,
        ["Requisito", "Estado"],
        [
            ["Mapa del sitio", "Cumplido"],
            ["Pagina 404 personalizada", "Cumplido"],
            ["Menu completo del sitio", "Cumplido"],
            ["Busqueda de productos", "Cumplido"],
            ["Catalogo con imagen, precio y disponibilidad", "Cumplido"],
            ["Registro, login y recuperacion", "Cumplido"],
            ["Contacto y buzon", "Cumplido"],
            ["Captcha matematico", "Cumplido"],
            ["Backend con JSON", "Cumplido"],
        ],
    )

    add_heading(document, "Evidencia del backend")
    backend_rows = [
        [check["method"], check["path"], str(check["status"]), check["sample"]]
        for check in evidence["checks"]
    ]
    add_table(document, ["Metodo", "Ruta", "Codigo", "Resultado"], backend_rows)
    add_paragraph(document, evidence["build"])
    add_paragraph(document, evidence["typecheck"])

    add_heading(document, "Capturas del sitio")
    captures = [
        ("01-inicio.png", "Pagina de inicio"),
        ("02-productos.png", "Catalogo de productos"),
        ("03-contacto-buzon.png", "Contacto y buzon"),
        ("04-registro.png", "Registro de usuario"),
        ("05-login.png", "Inicio de sesion"),
        ("06-ayuda.png", "Ayuda y preguntas frecuentes"),
        ("07-mapa-sitio.png", "Mapa del sitio"),
        ("08-error-404.png", "Pagina 404"),
    ]

    for filename, caption in captures:
        image_path = CAPTURES_DIR / filename
        if not image_path.exists():
            continue
        add_heading(document, caption, size=12)
        document.add_picture(str(image_path), width=Inches(6.1))
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.italic = True
        caption_run.font.name = "Arial"
        caption_run.font.size = Pt(10)

    document.add_page_break()

    add_heading(document, "Conclusiones")
    add_paragraph(
        document,
        "Como conclusion, considero que el proyecto de Ferreteria Valdez cumplio de buena manera con los objetivos planteados para esta unidad, ya que no solo se logro hacer una pagina con buena presentacion, sino tambien un sitio mas completo y funcional. A lo largo del desarrollo se integraron elementos importantes que ayudan a que la experiencia del usuario sea mejor, por ejemplo el catalogo con filtros, el formulario de contacto, la ayuda, el mapa del sitio y la pagina 404 personalizada. Tambien fue importante agregar validaciones en los formularios y una verificacion humana sencilla, porque eso hace que el sitio se vea mas serio y mejor pensado. Otra parte valiosa fue conectar el frontend con un backend basico, ya que asi se pudo guardar informacion de usuarios y mensajes en archivos JSON.\n\nEste trabajo tambien me dejo como aprendizaje que hacer una pagina web no solo es poner colores o texto en pantalla, sino organizar bien la informacion, pensar en lo que necesita el usuario y comprobar que todo funcione correctamente. Igualmente, me sirvio para practicar el uso de herramientas como GitHub, la estructura de un proyecto web y la importancia de documentar el proceso. En mi opinion, fue una actividad muy completa porque permitio aplicar conocimientos tecnicos y al mismo tiempo mejorar la manera de presentar un proyecto final.",
    )

    add_heading(document, "Repositorio")
    add_paragraph(document, REPO_URL)


def main() -> None:
    document = Document()
    set_page_layout(document)
    style_normal(document)
    build_cover(document)
    build_body(document)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
